
import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import os
from docx import Document

# Directly configure the API key
api_key = "AIzaSyBaNyjY1c6NqYc8JwR82NkUrGRjg5pO5bY"
genai.configure(api_key=api_key)

def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=api_key)
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    if not os.path.exists("faiss_index"):
        os.makedirs("faiss_index")
    vector_store.save_local("faiss_index")

def get_translation_chain(target_language):
    prompt_template = f"""
    Translate the following text to {target_language}. Make sure to keep the meaning and context intact. Do not provide any additional information.\n\n
    Text:\n {{text}}\n

    Translation:
    """

    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3, google_api_key=api_key)

    prompt = PromptTemplate(template=prompt_template, input_variables=["text"])
    chain = LLMChain(llm=model, prompt=prompt)

    return chain

def translate_document(target_language):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=api_key)
    
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search("Translate the document")

    chain = get_translation_chain(target_language)

    translations = []
    for doc in docs:
        response = chain({"text": doc.page_content}, return_only_outputs=True)
        
        if "text" in response:
            translations.append(response["text"])
        else:
            st.error(f"Unexpected response format: {response}")

    translated_text = "\n\n".join(translations)
    return translated_text

def save_translation_as_docx(translated_text, filename="translated_document.docx"):
    doc = Document()
    doc.add_heading('Translated Document', 0)
    doc.add_paragraph(translated_text)
    doc.save(filename)
    return filename

def main():
    st.set_page_config(page_title="Translate PDF", layout="wide")
    st.header("Translate PDF ")

    # Language selection
    languages = {
        "Hindi": "Hindi",
        "Marathi": "Marathi",
        "Tamil": "Tamil",
        "Telugu": "Telugu",
        "Kannada": "Kannada",
        "Malayalam": "Malayalam",
        "Arabic": "Arabic"
    }
    target_language = st.selectbox("Select the target language for translation:", list(languages.keys()))

    # Placeholder for translated text
    translated_text = st.empty()

    # Upload and process PDF
    with st.sidebar:
        st.title("Menu:")
        pdf_docs = st.file_uploader("Upload your PDF Files and Click on the Submit & Process Button", accept_multiple_files=True)
        if st.button("Submit & Process"):
            if pdf_docs:
                with st.spinner("Processing..."):
                    raw_text = get_pdf_text(pdf_docs)
                    text_chunks = get_text_chunks(raw_text)
                    get_vector_store(text_chunks)
                    st.success("Processing complete!")

    # Translate Document
    if st.button("Translate Document"):
        if not pdf_docs:
            st.warning("Please upload PDF files before translating.")
        else:
            with st.spinner("Translating..."):
                translated_text_value = translate_document(languages[target_language])
                if translated_text_value:
                    translated_text.text(translated_text_value)
                    st.success("Translation completed!")

                    # Save and download translated document
                    docx_filename = save_translation_as_docx(translated_text_value)
                    with open(docx_filename, "rb") as file:
                        st.download_button(
                            label="Download Translated Document (.docx)",
                            data=file,
                            file_name=docx_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

if __name__ == "__main__":
    main()
